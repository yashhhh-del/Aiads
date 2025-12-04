"""
AI Sales Copy & Ad Content Agent
================================
A complete AI system that generates multi-platform marketing content including:
- Google Ad headlines & descriptions
- Facebook/Instagram ad copies
- Captions + Hashtags
- SEO titles & meta descriptions
- CTA suggestions
- Keyword extraction
- Landing page content

Author: AI Sales Agent
Version: 2.0.0 (Groq FREE Edition)
"""

import streamlit as st
from groq import Groq
import os
import json
import sqlite3
import re
import io
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
import pandas as pd
import nltk
from collections import Counter

# Load environment variables
load_dotenv()

# =============================================================================
# DATABASE SETUP
# =============================================================================

def init_database():
    """Initialize SQLite database with required tables"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    # Users table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            email TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    
    # Content history table
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS content_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            business_name TEXT,
            business_type TEXT,
            product_service TEXT,
            target_audience TEXT,
            offer TEXT,
            tone TEXT,
            platform TEXT,
            headlines TEXT,
            descriptions TEXT,
            hashtags TEXT,
            keywords TEXT,
            cta TEXT,
            seo_title TEXT,
            meta_description TEXT,
            landing_page_content TEXT,
            full_response TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users (id)
        )
    ''')
    
    conn.commit()
    conn.close()

def save_to_history(user_id, inputs, outputs):
    """Save generated content to database history"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        INSERT INTO content_history 
        (user_id, business_name, business_type, product_service, target_audience, 
         offer, tone, platform, headlines, descriptions, hashtags, keywords, 
         cta, seo_title, meta_description, landing_page_content, full_response)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        user_id,
        inputs.get('business_name', ''),
        inputs.get('business_type', ''),
        inputs.get('product_service', ''),
        inputs.get('target_audience', ''),
        inputs.get('offer', ''),
        inputs.get('tone', ''),
        inputs.get('platform', ''),
        outputs.get('headlines', ''),
        outputs.get('descriptions', ''),
        outputs.get('hashtags', ''),
        outputs.get('keywords', ''),
        outputs.get('cta', ''),
        outputs.get('seo_title', ''),
        outputs.get('meta_description', ''),
        outputs.get('landing_page_content', ''),
        json.dumps(outputs)
    ))
    
    conn.commit()
    conn.close()

def get_user_history(user_id, limit=50):
    """Retrieve user's content generation history"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT * FROM content_history 
        WHERE user_id = ? 
        ORDER BY created_at DESC 
        LIMIT ?
    ''', (user_id, limit))
    
    columns = [description[0] for description in cursor.description]
    rows = cursor.fetchall()
    conn.close()
    
    return [dict(zip(columns, row)) for row in rows]

def authenticate_user(username, password):
    """Simple user authentication"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    cursor.execute('''
        SELECT id, username FROM users 
        WHERE username = ? AND password = ?
    ''', (username, password))
    
    user = cursor.fetchone()
    conn.close()
    
    return user

def register_user(username, password, email):
    """Register a new user"""
    conn = sqlite3.connect('sales_content.db')
    cursor = conn.cursor()
    
    try:
        cursor.execute('''
            INSERT INTO users (username, password, email) 
            VALUES (?, ?, ?)
        ''', (username, password, email))
        conn.commit()
        user_id = cursor.lastrowid
        conn.close()
        return user_id
    except sqlite3.IntegrityError:
        conn.close()
        return None

# =============================================================================
# NLP KEYWORD EXTRACTION ENGINE
# =============================================================================

def download_nltk_data():
    """Download required NLTK data"""
    nltk_packages = [
        ('tokenizers/punkt', 'punkt'),
        ('tokenizers/punkt_tab', 'punkt_tab'),
        ('corpora/stopwords', 'stopwords'),
        ('taggers/averaged_perceptron_tagger', 'averaged_perceptron_tagger'),
        ('taggers/averaged_perceptron_tagger_eng', 'averaged_perceptron_tagger_eng'),
    ]
    
    for path, package in nltk_packages:
        try:
            nltk.data.find(path)
        except LookupError:
            nltk.download(package, quiet=True)

def extract_keywords_nlp(text, num_keywords=15):
    """
    Extract keywords from text using NLTK
    Returns most relevant keywords based on frequency and POS tagging
    """
    download_nltk_data()
    
    from nltk.corpus import stopwords
    from nltk.tokenize import word_tokenize
    from nltk import pos_tag
    
    # Tokenize and clean
    tokens = word_tokenize(text.lower())
    stop_words = set(stopwords.words('english'))
    
    # Additional marketing stopwords
    marketing_stopwords = {
        'will', 'can', 'get', 'make', 'use', 'new', 'one', 'also', 
        'like', 'just', 'know', 'take', 'come', 'see', 'want', 'look',
        'give', 'think', 'good', 'best', 'way', 'need', 'feel', 'try'
    }
    stop_words.update(marketing_stopwords)
    
    # Filter tokens
    filtered_tokens = [
        token for token in tokens 
        if token.isalnum() and token not in stop_words and len(token) > 2
    ]
    
    # POS tagging - keep nouns, verbs, adjectives
    pos_tags = pos_tag(filtered_tokens)
    important_tags = {'NN', 'NNS', 'NNP', 'NNPS', 'VB', 'VBG', 'JJ', 'JJR', 'JJS'}
    
    important_words = [
        word for word, tag in pos_tags 
        if tag in important_tags
    ]
    
    # Get frequency distribution
    word_freq = Counter(important_words)
    keywords = [word for word, count in word_freq.most_common(num_keywords)]
    
    return keywords

def generate_hashtags(keywords, platform='instagram'):
    """Generate platform-appropriate hashtags from keywords"""
    hashtags = []
    
    for keyword in keywords[:10]:
        # Clean and format keyword
        tag = keyword.replace(' ', '').replace('-', '').lower()
        if tag:
            hashtags.append(f"#{tag}")
    
    # Add common marketing hashtags based on platform
    if platform.lower() in ['instagram', 'facebook']:
        common_tags = ['#marketing', '#business', '#entrepreneur', '#success', '#growth']
        hashtags.extend(common_tags[:3])
    
    return list(set(hashtags))[:15]

# =============================================================================
# PROMPT TEMPLATES ENGINE
# =============================================================================

class PromptTemplates:
    """Prompt templates for different platforms and content types"""
    
    @staticmethod
    def get_tone_modifier(tone):
        """Return tone-specific writing instructions"""
        tone_modifiers = {
            'Professional': "Use formal, business-appropriate language. Be authoritative and trustworthy. Focus on value propositions and credibility.",
            'Emotional': "Connect emotionally with the reader. Use storytelling elements. Appeal to feelings, desires, and aspirations.",
            'Exciting': "Use energetic, dynamic language. Create enthusiasm and anticipation. Use action words and exclamation points sparingly but effectively.",
            'Urgent': "Create a sense of urgency and scarcity. Use time-sensitive language. Emphasize limited availability or time-bound offers.",
            'Friendly': "Use warm, conversational tone. Be approachable and relatable. Write as if talking to a friend.",
            'Luxury': "Use sophisticated, premium language. Emphasize exclusivity and quality. Appeal to aspirational desires."
        }
        return tone_modifiers.get(tone, tone_modifiers['Professional'])
    
    @staticmethod
    def google_ads_prompt(inputs):
        """Generate Google Ads content prompt"""
        return f"""
You are an expert Google Ads copywriter. Create high-converting Google Ads content.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

STRICT CHARACTER LIMITS:
- Headlines: Maximum 30 characters each (including spaces)
- Descriptions: Maximum 90 characters each (including spaces)

Generate the following in JSON format:
{{
    "headlines": [
        // 10 unique headlines, each MUST be 30 characters or less
    ],
    "descriptions": [
        // 5 unique descriptions, each MUST be 90 characters or less
    ],
    "display_urls": [
        // 3 suggested display URL paths
    ],
    "keywords": [
        // 15 relevant search keywords for this ad
    ],
    "negative_keywords": [
        // 5 negative keywords to exclude
    ],
    "cta_suggestions": [
        // 5 call-to-action phrases
    ]
}}

IMPORTANT: 
- Count characters carefully - headlines over 30 chars will be rejected
- Make headlines compelling and action-oriented
- Include the main keyword in at least 3 headlines
- Descriptions should expand on the value proposition
- Return ONLY valid JSON, no extra text
"""

    @staticmethod
    def facebook_instagram_prompt(inputs):
        """Generate Facebook/Instagram ad content prompt"""
        return f"""
You are a social media marketing expert specializing in Facebook and Instagram ads.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

Generate the following in JSON format:
{{
    "facebook_ad": {{
        "primary_text": [
            // 3 variations of primary text (125 characters ideal, max 500)
        ],
        "headlines": [
            // 5 headlines (max 40 characters)
        ],
        "descriptions": [
            // 3 link descriptions (max 30 characters)
        ],
        "cta_button": [
            // 3 CTA button suggestions (Shop Now, Learn More, Sign Up, etc.)
        ]
    }},
    "instagram_ad": {{
        "captions": [
            // 3 engaging captions (optimal 138-150 characters, can go up to 2200)
        ],
        "story_text": [
            // 3 short story overlay texts (max 100 characters)
        ],
        "hashtags": [
            // 20 relevant hashtags
        ],
        "bio_link_cta": [
            // 3 "Link in bio" style CTAs
        ]
    }},
    "carousel_hooks": [
        // 5 carousel slide headline hooks
    ],
    "engagement_questions": [
        // 3 questions to boost engagement in comments
    ]
}}

IMPORTANT:
- Facebook primary text should hook in first 125 characters
- Instagram captions should be engaging and include emojis where appropriate
- Hashtags should mix popular and niche tags
- Include a strong hook in the first line
- Return ONLY valid JSON, no extra text
"""

    @staticmethod
    def seo_content_prompt(inputs):
        """Generate SEO-optimized content prompt"""
        return f"""
You are an SEO specialist and content strategist.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

Generate the following in JSON format:
{{
    "seo_titles": [
        // 5 SEO-optimized titles (50-60 characters)
    ],
    "meta_descriptions": [
        // 5 meta descriptions (150-160 characters)
    ],
    "h1_headings": [
        // 3 H1 heading suggestions
    ],
    "h2_subheadings": [
        // 5 H2 subheading suggestions
    ],
    "primary_keywords": [
        // 5 primary target keywords
    ],
    "secondary_keywords": [
        // 10 secondary/LSI keywords
    ],
    "long_tail_keywords": [
        // 10 long-tail keyword phrases
    ],
    "url_slugs": [
        // 3 SEO-friendly URL slug suggestions
    ],
    "image_alt_texts": [
        // 5 image alt text suggestions
    ],
    "schema_suggestions": {{
        "type": "suggested schema markup type",
        "key_properties": ["list of key schema properties to include"]
    }}
}}

IMPORTANT:
- Titles should include primary keyword near the beginning
- Meta descriptions should be compelling and include a CTA
- Keywords should have commercial/transactional intent
- Return ONLY valid JSON, no extra text
"""

    @staticmethod
    def landing_page_prompt(inputs):
        """Generate landing page content prompt"""
        return f"""
You are a conversion rate optimization expert and landing page copywriter.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

Generate the following in JSON format:
{{
    "hero_section": {{
        "headline": "Main headline (max 10 words)",
        "subheadline": "Supporting subheadline (max 20 words)",
        "cta_button_text": "Primary CTA button text",
        "cta_supporting_text": "Text below CTA (e.g., 'No credit card required')"
    }},
    "value_propositions": [
        {{
            "title": "Value prop title",
            "description": "2-3 sentence description",
            "icon_suggestion": "suggested icon name"
        }}
    ],
    "features_benefits": [
        {{
            "feature": "Feature name",
            "benefit": "How it benefits the user"
        }}
    ],
    "social_proof": {{
        "testimonial_prompts": [
            // 3 prompts for collecting testimonials
        ],
        "stats_suggestions": [
            // 3 types of stats to showcase
        ],
        "trust_badges": [
            // 5 trust badge suggestions
        ]
    }},
    "faq_questions": [
        {{
            "question": "FAQ question",
            "answer": "Concise answer"
        }}
    ],
    "urgency_elements": [
        // 3 urgency/scarcity elements
    ],
    "final_cta": {{
        "headline": "Final section headline",
        "cta_text": "Final CTA button text",
        "guarantee": "Risk reversal statement"
    }}
}}

IMPORTANT:
- Hero headline should be benefit-focused
- Value props should address pain points
- Include specific numbers where possible
- CTAs should be action-oriented
- Return ONLY valid JSON, no extra text
"""

    @staticmethod
    def multi_platform_prompt(inputs):
        """Generate content for all platforms at once"""
        return f"""
You are a multi-channel marketing strategist and copywriter.

BUSINESS DETAILS:
- Business Name: {inputs['business_name']}
- Business Type: {inputs['business_type']}
- Product/Service: {inputs['product_service']}
- Target Audience: {inputs['target_audience']}
- Offer: {inputs['offer']}
- Tone: {inputs['tone']}

TONE INSTRUCTIONS: {PromptTemplates.get_tone_modifier(inputs['tone'])}

Generate comprehensive marketing content for ALL platforms in JSON format:
{{
    "google_ads": {{
        "headlines": ["10 headlines, max 30 chars each"],
        "descriptions": ["5 descriptions, max 90 chars each"],
        "keywords": ["15 target keywords"]
    }},
    "facebook": {{
        "primary_texts": ["3 ad texts"],
        "headlines": ["5 headlines, max 40 chars"],
        "cta_buttons": ["3 CTA suggestions"]
    }},
    "instagram": {{
        "captions": ["3 engaging captions with emojis"],
        "hashtags": ["20 relevant hashtags"],
        "story_texts": ["3 story overlay texts"]
    }},
    "seo": {{
        "titles": ["5 SEO titles, 50-60 chars"],
        "meta_descriptions": ["5 meta descriptions, 150-160 chars"],
        "keywords": {{
            "primary": ["5 primary keywords"],
            "secondary": ["10 secondary keywords"],
            "long_tail": ["10 long-tail phrases"]
        }}
    }},
    "landing_page": {{
        "hero_headline": "Main headline",
        "hero_subheadline": "Supporting text",
        "value_props": ["4 value propositions"],
        "cta_texts": ["3 CTA variations"],
        "testimonial_prompts": ["3 testimonial collection prompts"]
    }},
    "email": {{
        "subject_lines": ["5 email subject lines"],
        "preview_texts": ["3 preview/preheader texts"],
        "cta_buttons": ["3 email CTA texts"]
    }},
    "general": {{
        "taglines": ["5 brand taglines"],
        "elevator_pitch": "30-second elevator pitch",
        "unique_selling_points": ["3 USPs"]
    }}
}}

Ensure all content is cohesive across platforms while optimized for each platform's best practices.
Return ONLY valid JSON, no extra text or markdown.
"""

# =============================================================================
# LLM CONTENT GENERATION ENGINE (GROQ - FREE)
# =============================================================================

class ContentGenerator:
    """Main content generation engine using Groq (FREE & FAST)"""
    
    def __init__(self, api_key):
        self.client = Groq(api_key=api_key)
        self.model = "llama-3.3-70b-versatile"  # Free and powerful
    
    def generate_content(self, prompt, max_tokens=4000):
        """Generate content using Groq API"""
        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system",
                        "content": "You are an expert marketing copywriter. Always respond with valid JSON only. No markdown, no code blocks, no explanations - just pure JSON that can be parsed directly."
                    },
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                max_tokens=max_tokens,
                temperature=0.7
            )
            
            content = response.choices[0].message.content.strip()
            
            # Clean JSON response
            if content.startswith('```'):
                content = re.sub(r'^```json?\n?', '', content)
                content = re.sub(r'\n?```$', '', content)
            
            # Try to find JSON in the response
            json_match = re.search(r'\{[\s\S]*\}', content)
            if json_match:
                content = json_match.group()
            
            return json.loads(content)
            
        except json.JSONDecodeError as e:
            st.error(f"Error parsing response: {e}")
            st.code(content[:500] if 'content' in dir() else "No content")
            return None
        except Exception as e:
            st.error(f"API Error: {e}")
            return None
    
    def generate_google_ads(self, inputs):
        """Generate Google Ads content"""
        prompt = PromptTemplates.google_ads_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_social_media(self, inputs):
        """Generate Facebook/Instagram content"""
        prompt = PromptTemplates.facebook_instagram_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_seo_content(self, inputs):
        """Generate SEO-optimized content"""
        prompt = PromptTemplates.seo_content_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_landing_page(self, inputs):
        """Generate landing page content"""
        prompt = PromptTemplates.landing_page_prompt(inputs)
        return self.generate_content(prompt)
    
    def generate_all_platforms(self, inputs):
        """Generate content for all platforms"""
        prompt = PromptTemplates.multi_platform_prompt(inputs)
        return self.generate_content(prompt, max_tokens=6000)

# =============================================================================
# EXPORT FUNCTIONS (PDF & DOCX)
# =============================================================================

def export_to_docx(content_data, inputs):
    """Export generated content to Word document"""
    doc = Document()
    
    # Title
    title = doc.add_heading('AI Generated Marketing Content', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Business Info Section
    doc.add_heading('Business Information', level=1)
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_items = [
        ('Business Name', inputs.get('business_name', '')),
        ('Business Type', inputs.get('business_type', '')),
        ('Product/Service', inputs.get('product_service', '')),
        ('Target Audience', inputs.get('target_audience', '')),
        ('Offer', inputs.get('offer', '')),
        ('Tone', inputs.get('tone', ''))
    ]
    
    for i, (label, value) in enumerate(info_items):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    # Add content sections based on available data
    def add_list_section(title, items):
        if items:
            doc.add_heading(title, level=2)
            for item in items:
                doc.add_paragraph(f"• {item}", style='List Bullet')
            doc.add_paragraph()
    
    def add_dict_section(title, data, level=2):
        if data:
            doc.add_heading(title, level=level)
            if isinstance(data, dict):
                for key, value in data.items():
                    if isinstance(value, list):
                        doc.add_heading(key.replace('_', ' ').title(), level=level+1)
                        for item in value:
                            if isinstance(item, dict):
                                for k, v in item.items():
                                    doc.add_paragraph(f"{k}: {v}")
                            else:
                                doc.add_paragraph(f"• {item}", style='List Bullet')
                    elif isinstance(value, dict):
                        add_dict_section(key.replace('_', ' ').title(), value, level+1)
                    else:
                        doc.add_paragraph(f"{key.replace('_', ' ').title()}: {value}")
            doc.add_paragraph()
    
    # Process content data
    if isinstance(content_data, dict):
        for section_name, section_content in content_data.items():
            section_title = section_name.replace('_', ' ').title()
            
            if isinstance(section_content, list):
                add_list_section(section_title, section_content)
            elif isinstance(section_content, dict):
                add_dict_section(section_title, section_content)
            else:
                doc.add_heading(section_title, level=2)
                doc.add_paragraph(str(section_content))
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save to bytes
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer

def export_to_pdf(content_data, inputs):
    """Export generated content to PDF"""
    pdf_buffer = io.BytesIO()
    doc = SimpleDocTemplate(pdf_buffer, pagesize=A4,
                           rightMargin=72, leftMargin=72,
                           topMargin=72, bottomMargin=18)
    
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=14,
        spaceBefore=20,
        spaceAfter=10,
        textColor=colors.darkblue
    )
    
    subheading_style = ParagraphStyle(
        'CustomSubheading',
        parent=styles['Heading3'],
        fontSize=12,
        spaceBefore=15,
        spaceAfter=8
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=10,
        spaceBefore=5,
        spaceAfter=5
    )
    
    bullet_style = ParagraphStyle(
        'CustomBullet',
        parent=styles['Normal'],
        fontSize=10,
        leftIndent=20,
        spaceBefore=3,
        spaceAfter=3
    )
    
    story = []
    
    # Title
    story.append(Paragraph("AI Generated Marketing Content", title_style))
    story.append(Spacer(1, 20))
    
    # Business Info
    story.append(Paragraph("Business Information", heading_style))
    
    info_data = [
        ['Field', 'Value'],
        ['Business Name', inputs.get('business_name', '')],
        ['Business Type', inputs.get('business_type', '')],
        ['Product/Service', inputs.get('product_service', '')],
        ['Target Audience', inputs.get('target_audience', '')],
        ['Offer', inputs.get('offer', '')],
        ['Tone', inputs.get('tone', '')]
    ]
    
    info_table = Table(info_data, colWidths=[2*inch, 4*inch])
    info_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black)
    ]))
    
    story.append(info_table)
    story.append(Spacer(1, 20))
    
    def add_content_to_story(data, level=0):
        """Recursively add content to PDF story"""
        if isinstance(data, dict):
            for key, value in data.items():
                title = key.replace('_', ' ').title()
                
                if level == 0:
                    story.append(Paragraph(title, heading_style))
                else:
                    story.append(Paragraph(title, subheading_style))
                
                if isinstance(value, list):
                    for item in value:
                        if isinstance(item, dict):
                            for k, v in item.items():
                                # Escape special characters for PDF
                                text = str(v).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                                story.append(Paragraph(f"<b>{k}:</b> {text}", bullet_style))
                        else:
                            text = str(item).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            story.append(Paragraph(f"• {text}", bullet_style))
                elif isinstance(value, dict):
                    add_content_to_story(value, level + 1)
                else:
                    text = str(value).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(text, body_style))
                
                story.append(Spacer(1, 10))
        elif isinstance(data, list):
            for item in data:
                text = str(item).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                story.append(Paragraph(f"• {text}", bullet_style))
    
    # Add content
    if content_data:
        add_content_to_story(content_data)
    
    # Footer
    story.append(Spacer(1, 30))
    story.append(Paragraph(
        f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        ParagraphStyle('Footer', parent=styles['Normal'], alignment=1, fontSize=8)
    ))
    
    # Build PDF
    doc.build(story)
    pdf_buffer.seek(0)
    
    return pdf_buffer

# =============================================================================
# STREAMLIT UI COMPONENTS
# =============================================================================

def render_sidebar():
    """Render the sidebar with navigation and settings"""
    with st.sidebar:
        st.image("https://img.icons8.com/3d-fluency/94/artificial-intelligence.png", width=80)
        st.title("🚀 AI Sales Agent")
        st.markdown("**FREE Edition (Groq)**")
        st.markdown("---")
        
        # Navigation
        page = st.radio(
            "Navigation",
            ["🏠 Home", "✨ Generate Content", "📊 Dashboard", "⚙️ Settings"],
            label_visibility="collapsed"
        )
        
        st.markdown("---")
        
        # API Key input
        api_key = st.text_input(
            "Groq API Key (FREE)",
            type="password",
            value=st.session_state.get('api_key', ''),
            help="Get FREE API key from console.groq.com"
        )
        
        if api_key:
            st.session_state['api_key'] = api_key
            st.success("✅ API Key configured")
        else:
            st.warning("⚠️ Enter Groq API Key")
            st.markdown("[Get FREE Key →](https://console.groq.com)")
        
        st.markdown("---")
        st.markdown("### 📌 Quick Tips")
        st.markdown("""
        - Be specific about your target audience
        - Include your unique selling proposition
        - Mention any time-sensitive offers
        - Choose the right tone for your brand
        """)
        
        st.markdown("---")
        st.markdown("### 🆓 Why Groq?")
        st.markdown("""
        - ✅ **100% FREE**
        - ✅ Super fast (10x faster)
        - ✅ No credit card needed
        - ✅ Llama 3.3 70B model
        """)
        
        return page

def render_input_form():
    """Render the main input form"""
    st.subheader("📝 Business Details")
    
    col1, col2 = st.columns(2)
    
    with col1:
        business_name = st.text_input(
            "Business Name *",
            placeholder="e.g., TechFlow Solutions"
        )
        
        business_type = st.selectbox(
            "Business Type *",
            [
                "E-commerce",
                "SaaS",
                "Local Service",
                "Consulting",
                "Healthcare",
                "Education",
                "Real Estate",
                "Finance",
                "Food & Restaurant",
                "Fitness & Wellness",
                "Travel & Tourism",
                "Manufacturing",
                "Other"
            ]
        )
        
        product_service = st.text_area(
            "Product/Service Description *",
            placeholder="Describe your product or service in detail...",
            height=100
        )
    
    with col2:
        target_audience = st.text_area(
            "Target Audience *",
            placeholder="e.g., Small business owners aged 25-45, tech-savvy, looking to automate their workflow...",
            height=80
        )
        
        offer = st.text_input(
            "Current Offer/Promotion",
            placeholder="e.g., 50% off for first 100 customers"
        )
        
        tone = st.selectbox(
            "Content Tone *",
            ["Professional", "Emotional", "Exciting", "Urgent", "Friendly", "Luxury"]
        )
    
    st.markdown("---")
    
    platform = st.multiselect(
        "Select Platforms *",
        [
            "All Platforms",
            "Google Ads",
            "Facebook",
            "Instagram",
            "SEO Content",
            "Landing Page"
        ],
        default=["All Platforms"]
    )
    
    return {
        'business_name': business_name,
        'business_type': business_type,
        'product_service': product_service,
        'target_audience': target_audience,
        'offer': offer,
        'tone': tone,
        'platform': platform
    }

def display_content_results(results, platform_type):
    """Display generated content results"""
    if not results:
        st.warning("No content generated. Please try again.")
        return
    
    st.success("✅ Content generated successfully!")
    
    # Display based on content structure
    if isinstance(results, dict):
        for section_name, section_content in results.items():
            with st.expander(f"📌 {section_name.replace('_', ' ').title()}", expanded=True):
                if isinstance(section_content, list):
                    for i, item in enumerate(section_content, 1):
                        if isinstance(item, dict):
                            for key, value in item.items():
                                st.markdown(f"**{key}:** {value}")
                            st.markdown("---")
                        else:
                            st.markdown(f"{i}. {item}")
                elif isinstance(section_content, dict):
                    for key, value in section_content.items():
                        st.markdown(f"**{key.replace('_', ' ').title()}:**")
                        if isinstance(value, list):
                            for item in value:
                                if isinstance(item, dict):
                                    for k, v in item.items():
                                        st.markdown(f"  - **{k}:** {v}")
                                else:
                                    st.markdown(f"  - {item}")
                        elif isinstance(value, dict):
                            for k, v in value.items():
                                st.markdown(f"  - **{k}:** {v}")
                        else:
                            st.markdown(f"  {value}")
                else:
                    st.write(section_content)

def render_generate_page():
    """Render the content generation page"""
    st.title("✨ Generate Marketing Content")
    st.markdown("Create high-converting marketing content powered by **FREE Groq AI**")
    
    # Check API key
    if 'api_key' not in st.session_state or not st.session_state['api_key']:
        st.warning("⚠️ Please enter your FREE Groq API key in the sidebar to continue.")
        st.info("👉 Get your FREE API key at [console.groq.com](https://console.groq.com)")
        return
    
    # Input form
    inputs = render_input_form()
    
    # Validate inputs
    required_fields = ['business_name', 'business_type', 'product_service', 'target_audience']
    is_valid = all(inputs.get(field) for field in required_fields)
    
    col1, col2, col3 = st.columns([1, 1, 1])
    
    with col2:
        generate_btn = st.button(
            "🚀 Generate Content",
            type="primary",
            use_container_width=True,
            disabled=not is_valid
        )
    
    if not is_valid and generate_btn:
        st.error("Please fill in all required fields (*)")
        return
    
    if generate_btn:
        with st.spinner("🔄 Generating content with Groq AI... This is fast!"):
            try:
                generator = ContentGenerator(st.session_state['api_key'])
                
                results = {}
                platforms = inputs['platform']
                
                if "All Platforms" in platforms:
                    results = generator.generate_all_platforms(inputs)
                else:
                    if "Google Ads" in platforms:
                        results['google_ads'] = generator.generate_google_ads(inputs)
                    if "Facebook" in platforms or "Instagram" in platforms:
                        social_results = generator.generate_social_media(inputs)
                        if social_results:
                            results.update(social_results)
                    if "SEO Content" in platforms:
                        results['seo'] = generator.generate_seo_content(inputs)
                    if "Landing Page" in platforms:
                        results['landing_page'] = generator.generate_landing_page(inputs)
                
                if results:
                    # Store in session state
                    st.session_state['last_results'] = results
                    st.session_state['last_inputs'] = inputs
                    
                    # Extract keywords using NLP
                    text_for_nlp = f"{inputs['product_service']} {inputs['target_audience']} {inputs['offer']}"
                    nlp_keywords = extract_keywords_nlp(text_for_nlp)
                    
                    st.session_state['nlp_keywords'] = nlp_keywords
                    
                    # Save to history (user_id = 1 for demo)
                    # Convert platform list to string for database
                    inputs_for_db = inputs.copy()
                    inputs_for_db['platform'] = ', '.join(inputs['platform']) if isinstance(inputs['platform'], list) else inputs['platform']
                    
                    flat_outputs = {
                        'headlines': json.dumps(results.get('google_ads', {}).get('headlines', [])),
                        'descriptions': json.dumps(results.get('google_ads', {}).get('descriptions', [])),
                        'hashtags': json.dumps(results.get('instagram', {}).get('hashtags', [])),
                        'keywords': json.dumps(nlp_keywords),
                        'cta': json.dumps(results.get('google_ads', {}).get('cta_suggestions', [])),
                        'seo_title': json.dumps(results.get('seo', {}).get('titles', [])),
                        'meta_description': json.dumps(results.get('seo', {}).get('meta_descriptions', [])),
                        'landing_page_content': json.dumps(results.get('landing_page', {}))
                    }
                    save_to_history(1, inputs_for_db, flat_outputs)
                    
            except Exception as e:
                st.error(f"Error generating content: {str(e)}")
                return
        
        # Display results
        if 'last_results' in st.session_state:
            st.markdown("---")
            st.header("📄 Generated Content")
            
            # NLP Keywords section
            if 'nlp_keywords' in st.session_state:
                with st.expander("🔑 Extracted Keywords (NLP)", expanded=True):
                    keywords = st.session_state['nlp_keywords']
                    hashtags = generate_hashtags(keywords)
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.markdown("**Keywords:**")
                        st.write(", ".join(keywords))
                    with col2:
                        st.markdown("**Generated Hashtags:**")
                        st.write(" ".join(hashtags))
            
            display_content_results(st.session_state['last_results'], inputs['platform'])
            
            # Export buttons
            st.markdown("---")
            st.subheader("📥 Export Content")
            
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # DOCX Export
                docx_buffer = export_to_docx(
                    st.session_state['last_results'],
                    st.session_state['last_inputs']
                )
                st.download_button(
                    label="📄 Download DOCX",
                    data=docx_buffer,
                    file_name=f"marketing_content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            with col2:
                # PDF Export
                pdf_buffer = export_to_pdf(
                    st.session_state['last_results'],
                    st.session_state['last_inputs']
                )
                st.download_button(
                    label="📕 Download PDF",
                    data=pdf_buffer,
                    file_name=f"marketing_content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                    mime="application/pdf"
                )
            
            with col3:
                # JSON Export
                json_str = json.dumps(st.session_state['last_results'], indent=2)
                st.download_button(
                    label="📋 Download JSON",
                    data=json_str,
                    file_name=f"marketing_content_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json"
                )

def render_dashboard():
    """Render the dashboard with history"""
    st.title("📊 Content Dashboard")
    st.markdown("View and manage your generated content history")
    
    # Get history
    history = get_user_history(1, limit=20)
    
    if not history:
        st.info("No content history yet. Generate some content to see it here!")
        return
    
    # Summary metrics
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total Generated", len(history))
    
    with col2:
        platforms = [h['platform'] for h in history if h['platform']]
        st.metric("Platforms Used", len(set(platforms)))
    
    with col3:
        tones = [h['tone'] for h in history if h['tone']]
        most_common_tone = max(set(tones), key=tones.count) if tones else "N/A"
        st.metric("Most Used Tone", most_common_tone)
    
    with col4:
        today_count = len([h for h in history if h['created_at'] and datetime.now().strftime('%Y-%m-%d') in h['created_at']])
        st.metric("Generated Today", today_count)
    
    st.markdown("---")
    
    # History table
    st.subheader("📜 Recent Content")
    
    for i, record in enumerate(history[:10]):
        with st.expander(f"📌 {record['business_name']} - {record['created_at'][:10] if record['created_at'] else 'N/A'}"):
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"**Business Type:** {record['business_type']}")
                st.markdown(f"**Platform:** {record['platform']}")
                st.markdown(f"**Tone:** {record['tone']}")
            
            with col2:
                st.markdown(f"**Target Audience:** {record['target_audience']}")
                st.markdown(f"**Offer:** {record['offer']}")
            
            if record['full_response']:
                try:
                    content = json.loads(record['full_response'])
                    st.json(content)
                except:
                    st.text(record['full_response'])

def render_settings():
    """Render settings page"""
    st.title("⚙️ Settings")
    
    st.subheader("🔑 API Configuration")
    
    api_key = st.text_input(
        "Groq API Key (FREE)",
        type="password",
        value=st.session_state.get('api_key', ''),
        help="Get your FREE API key from console.groq.com"
    )
    
    if api_key:
        st.session_state['api_key'] = api_key
    
    st.markdown("👉 [Get FREE Groq API Key](https://console.groq.com)")
    
    model = st.selectbox(
        "Model",
        ["llama-3.3-70b-versatile", "llama-3.1-8b-instant", "mixtral-8x7b-32768"],
        help="Select the Groq model to use (all FREE!)"
    )
    
    st.session_state['model'] = model
    
    st.markdown("---")
    
    st.subheader("🎨 Content Preferences")
    
    default_tone = st.selectbox(
        "Default Tone",
        ["Professional", "Emotional", "Exciting", "Urgent", "Friendly", "Luxury"],
        index=0
    )
    
    st.session_state['default_tone'] = default_tone
    
    include_emojis = st.checkbox("Include emojis in social media content", value=True)
    st.session_state['include_emojis'] = include_emojis
    
    st.markdown("---")
    
    st.subheader("📊 Data Management")
    
    if st.button("🗑️ Clear All History", type="secondary"):
        conn = sqlite3.connect('sales_content.db')
        cursor = conn.cursor()
        cursor.execute("DELETE FROM content_history")
        conn.commit()
        conn.close()
        st.success("History cleared successfully!")
        st.rerun()

def render_home():
    """Render home page"""
    st.title("🚀 AI Sales Copy & Ad Content Agent")
    st.markdown("### Generate High-Converting Marketing Content in Seconds")
    st.markdown("#### 🆓 **100% FREE** - Powered by Groq AI")
    
    st.markdown("---")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.markdown("### 🎯 Google Ads")
        st.markdown("""
        - Headlines (30 chars)
        - Descriptions (90 chars)
        - Keywords targeting
        - Negative keywords
        """)
    
    with col2:
        st.markdown("### 📱 Social Media")
        st.markdown("""
        - Facebook ad copies
        - Instagram captions
        - Hashtag suggestions
        - Story content
        """)
    
    with col3:
        st.markdown("### 🔍 SEO Content")
        st.markdown("""
        - Meta titles
        - Meta descriptions
        - Keyword research
        - URL suggestions
        """)
    
    st.markdown("---")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 🏠 Landing Pages")
        st.markdown("""
        - Hero sections
        - Value propositions
        - CTAs
        - FAQ content
        """)
    
    with col2:
        st.markdown("### 📤 Export Options")
        st.markdown("""
        - PDF reports
        - Word documents
        - JSON data
        - Copy to clipboard
        """)
    
    st.markdown("---")
    
    # Getting started box
    st.info("""
    **🚀 Getting Started (3 Easy Steps):**
    1. Get FREE API key from [console.groq.com](https://console.groq.com)
    2. Enter API key in sidebar 👈
    3. Navigate to 'Generate Content' and start creating!
    """)
    
    # Quick demo section
    st.markdown("### 🎬 How It Works")
    
    steps = [
        ("1️⃣", "Enter Business Details", "Provide your business name, type, product/service, and target audience"),
        ("2️⃣", "Select Platforms", "Choose which platforms you want content for"),
        ("3️⃣", "Generate Content", "Our AI creates optimized content for each platform"),
        ("4️⃣", "Export & Use", "Download as PDF, DOCX, or copy directly")
    ]
    
    cols = st.columns(4)
    for i, (icon, title, desc) in enumerate(steps):
        with cols[i]:
            st.markdown(f"### {icon}")
            st.markdown(f"**{title}**")
            st.markdown(desc)

# =============================================================================
# MAIN APPLICATION
# =============================================================================

def main():
    """Main application entry point"""
    # Page config
    st.set_page_config(
        page_title="AI Sales Copy Agent (FREE)",
        page_icon="🚀",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Custom CSS
    st.markdown("""
    <style>
    .stButton>button {
        width: 100%;
        border-radius: 10px;
        height: 50px;
        font-weight: bold;
    }
    .stTextInput>div>div>input {
        border-radius: 10px;
    }
    .stTextArea>div>div>textarea {
        border-radius: 10px;
    }
    .stSelectbox>div>div>select {
        border-radius: 10px;
    }
    div[data-testid="stExpander"] div[role="button"] p {
        font-size: 1.1rem;
        font-weight: 600;
    }
    .main .block-container {
        padding-top: 2rem;
    }
    </style>
    """, unsafe_allow_html=True)
    
    # Initialize database
    init_database()
    
    # Initialize session state
    if 'api_key' not in st.session_state:
        st.session_state['api_key'] = os.getenv('GROQ_API_KEY', '')
    
    if 'last_results' not in st.session_state:
        st.session_state['last_results'] = None
    
    if 'last_inputs' not in st.session_state:
        st.session_state['last_inputs'] = None
    
    # Render sidebar and get current page
    page = render_sidebar()
    
    # Render appropriate page based on navigation
    if page == "🏠 Home":
        render_home()
    elif page == "✨ Generate Content":
        render_generate_page()
    elif page == "📊 Dashboard":
        render_dashboard()
    elif page == "⚙️ Settings":
        render_settings()


if __name__ == "__main__":
    main()
